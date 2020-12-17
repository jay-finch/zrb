# Parser  Imports
import pdfplumber
import pdf2image
import PIL
import pytesseract
import docx

# Spacy Imports
import spacy
from spacy.matcher import Matcher

# Utilities Imports
import os
import pandas as pd



#### Text Loading Function Library ####

def pdf_pdfplumber_textract(filepath_string):
    with pdfplumber.open(filepath_string) as pdf:
        resume_pages_text = [pdf.pages[i].extract_text() for i in range(len(pdf.pages))]
        resume_text = '\n'.join(resume_pages_text)
        return resume_text


def pdf_pdf2image_textract(filepath_string):
    images = pdf2image.convert_from_path(filepath_string)
    text = ''
    for img in images:
        img_text = pytesseract.image_to_string(img)
        text = ' '.join([text, img_text])
    return text


def pdf_textract(filepath_string, verbose=False):
    try:
        if verbose: print('Running the Plumber')
        return pdf_pdfplumber_textract(filepath_string)
    except:
        if verbose: print('Running the PDF2Image')
        return pdf_pdf2image_textract(filepath_string)


def docx_textract(filepath_string):
    doc = docx.Document(filepath_string)
    resume_text = '\n'.join([doc.paragraphs[i].text.strip() for i in range(len(doc.paragraphs))])
    return resume_text


def textract(filepath_string, verbose=False):
    if filepath_string.endswith('.docx'):
        if verbose: print(f'Running Function: docx_textract({filepath_string})')
        return docx_textract(filepath_string).replace('Jefferson Frank\n','')
    elif filepath_string.endswith('.pdf'):
        return pdf_textract(filepath_string, verbose=verbose).replace('Jefferson Frank\n','')
    else:
        return 'Text Not Extracted'


### Parsing Library ####

# Load Spacy Model and Matcher
nlp = spacy.load('en_core_web_sm') # load pretrained spacy model


def extract_name(resume_text):
    nlp_text = nlp(resume_text)
    matcher = Matcher(nlp.vocab) # load model's matcher

    pattern = [{'POS': 'PROPN'}, {'POS': 'PROPN'}]
    matcher.add('NAME', None, pattern)
    matches = matcher(nlp_text)

    for match_id, start, end in matches:
        span = nlp_text[start:end]
        return span.text


def extract_phone_number(resume_text):
    nlp_text = nlp(resume_text)
    matcher = Matcher(nlp.vocab)

    # Pattern 1
    # (000) 000-0000 
    p1 = [
        {"ORTH": "(", "OP": "?"},
        {"SHAPE": "ddd"}, 
        {"ORTH": ")", "OP": "*"},
        {"ORTH": "-", "OP": "*"},
        {"SHAPE": "ddd"}, 
        {"ORTH": "-", "OP": "*"}, 
        {"SHAPE": "dddd"}
        ]
    
    # Pattern 2
    # (000)000 - 0000
    p2 = [
        {"ORTH": "(", "OP": "?"},
        {"TEXT": {"REGEX": "^\d{3}\)*\-*\d{3}$"}}, 
        {"ORTH": "-", "OP": "?"}, 
        {"SHAPE": "dddd"}
        ]
    

    # Pattern 3
    # (000)000 - 0000
    p3 = [
        {"ORTH": "(", "OP": "?"},
        {"TEXT": {"REGEX": "^\d{3}\)*\-*\d{3}\-*\d{4}$"}}, 
        ]


    matcher.add("PhoneNumber", None, p1, p2, p3)
    matches = matcher(nlp_text)

    for match_id, start, end in matches:
        span = nlp_text[start:end]
        return span.text


def extract_email(resume_text):
    nlp_text = nlp(resume_text)
    matcher = Matcher(nlp.vocab)
    
    email_pattern = [{"LIKE_EMAIL": True}]

    matcher.add("Email", None, email_pattern)
    matches = matcher(nlp_text)

    for match_id, start, end in matches:
        span = nlp_text[start:end]
        return span.text


def find_section_header(section_dictionary, resume_text, verbose=False):
    if verbose: print(f'SYSTEM: Loading section patterns for {list(section_dictionary.keys())[0]}')
    section_patterns = list(section_dictionary.values())[0]
    
    # Create Spacy Objects: NLP object and matcher object
    nlp_text = nlp(resume_text)
    matcher = Matcher(nlp.vocab)

    # Add patterns to matcher object and find the matches
    matcher.add(
        list(section_dictionary.keys())[0], # name of the section
        None,
        *section_patterns,
        )
    matches = matcher(nlp_text)

    # Get string match
    if len(matches) < 0:
        return '*NO MATCH*'
    for match_id, start, end in matches:
        span = nlp_text[start:end]
        return span.text


def extract_resume_sections(resume_text):
    # Define Resume Section Paterns
    professional_summary = {'professional_summary': [
        ## Pattern 1
        [{'LOWER': 'professional', 'OP': '*'},
        {'LOWER': 'work', 'OP': '*'},
        {'LOWER': 'summary', 'OP': '+'},
        {'IS_PUNCT': True, 'OP': '*'}],
        ]}
    education = {'education':[
        ## Pattern 1
        [{'LOWER': 'education'},
        {'IS_PUNCT': True, 'OP': '*'}],
        ]}
    certifications = {'certifications':[
        ## Pattern 1
        [{'LOWER': 'certifications'},
        {'IS_PUNCT': True, 'OP': '*'}]
        ]}
    work_experience = {'work_experience': [
        ## Pattern 1
        [{'LOWER': 'professional', 'OP': '*'},
        {'LOWER': 'work', 'OP': '*'},
        {'LOWER': 'experience', 'IS_LOWER':False, 'OP': '+'},
        {'IS_PUNCT': True, 'OP': '*'},
        {'IS_LOWER': False,}],
        ## Pattern 2    
        [{'LOWER': 'professional', 'OP': '*'},
        {'LOWER': 'work', 'OP': '+'},
        {'LOWER': 'history', 'OP': '+'},
        {'IS_PUNCT': True, 'OP': '*'}],    
        ]}
    resume_section_patterns = [professional_summary, education, certifications, work_experience]
    

    # Seperate resume into key sections
    section_tuples_list = [(list(section_pattern.keys())[0], find_section_header(section_pattern, resume_text)) for section_pattern in resume_section_patterns]
    section_seperator='***************************'
    new_text = resume_text
    for name, string in section_tuples_list:
        try:
            new_text = new_text.replace(string,f'{section_seperator}\n{name}*')
        except:
            continue
    sections = new_text.split(section_seperator)

    # Buid a dictionary of resume sections
    def validated_section(section_title_string, sections=sections):
        section_list = [section.replace(f'{section_title_string}*','').strip() for section in sections if f'{section_title_string}*' in section]
        if len(section_list) == 0: 
            return ''
        return section_list[0]
    
    resume_sections = {
        'work_experience' : validated_section('work_experience'),
        'education' : validated_section('education'),
        'certifications' : validated_section('certifications'),
        'professional_summary' : validated_section('professional_summary'),
    }

    return resume_sections



### Master Parse Function ###

def parse_resume(filepath_string, verbose=False):
    if verbose: print(f'System: Extracting text from {filepath_string}')
    resume_text = textract(filepath_string)

    if verbose: print(f'System: Building resume object')
    key_resume_sections = extract_resume_sections(resume_text)
    resume_dictionary = {
        'name': extract_name(resume_text),
        'phone_number': extract_phone_number(resume_text),
        'email': extract_email(resume_text),
    }
    resume_dictionary.update(key_resume_sections)
    resume_dictionary['text'] = resume_text
    
    
    return resume_dictionary

####  Application Code ####
        
filepath_string = '/Users/jayfinch/Polydelta_Drive/3. Zimagi/Zimagi_Resume_Bot/Positions/SeniorSecurityEngineer/Resumes/The Maven Group Candidate-Eric Davidson 003.docx'

for key, value in parse_resume(filepath_string).items():
    print(f'\n{key}:\n{value}\n\n\n{"#"*200}')